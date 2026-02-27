#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
测试 Markdown 转 Word 时不会生成重复文字
"""

import os
import sys
import io
import tempfile

# 设置标准输出编码为 utf-8，避免 Windows cmd 的 GBK 编码问题
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

# 将项目根目录加入 Python 路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from src.core.markdown.markdown_converter import MarkdownConverter


def test_no_duplicate_text():
    """
    测试转换后的 Word 文档中不存在重复文字。
    使用与 bug/功能汇报.md 相同结构的内容进行测试。
    """
    md_content = """# skywaking功能汇报

## 一、功能概述

本系统基于SkyWalking构建了完整的链路观测能力，实现了业务系统全链路可观测性。

**核心亮点**：问题定位时间从**小时级降低至分钟级**，效率提升**90%以上**。

## 二、详细工作内容

### 2.1 功能开发

- **业务系统全景大屏**：基于ECharts实现业务系统健康度全景展示。

- **调用链拓扑可视化**：基于G6图可视化引擎实现业务系统全景拓扑展示。
"""

    # 创建临时文件
    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    temp_file.close()
    output_path = temp_file.name

    try:
        # 执行转换
        converter = MarkdownConverter()
        result = converter.to_word(md_content, output_path)
        assert result is True, "转换应返回 True"

        # 读取生成的 Word 文档
        doc = Document(output_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        print("=== 生成的段落内容 ===")
        for i, text in enumerate(paragraphs):
            print(f"  [{i}] {text}")

        # 验证关键文本只出现一次
        all_text = "\n".join(paragraphs)

        # 检查标题不重复
        assert (
            all_text.count("skywaking功能汇报") == 1
        ), f"'skywaking功能汇报' 出现了 {all_text.count('skywaking功能汇报')} 次，预期 1 次"

        assert (
            all_text.count("一、功能概述") == 1
        ), f"'一、功能概述' 出现了 {all_text.count('一、功能概述')} 次，预期 1 次"

        # 检查正文不重复
        assert (
            all_text.count("本系统基于SkyWalking构建了完整的链路观测能力") == 1
        ), f"正文段落出现重复"

        # 检查带粗体的段落不重复
        assert (
            all_text.count("核心亮点") == 1
        ), f"'核心亮点' 出现了 {all_text.count('核心亮点')} 次，预期 1 次"

        assert (
            all_text.count("小时级降低至分钟级") == 1
        ), f"'小时级降低至分钟级' 出现了 {all_text.count('小时级降低至分钟级')} 次，预期 1 次"

        # 检查列表项不重复
        assert (
            all_text.count("业务系统全景大屏") == 1
        ), f"'业务系统全景大屏' 出现了 {all_text.count('业务系统全景大屏')} 次，预期 1 次"

        assert (
            all_text.count("调用链拓扑可视化") == 1
        ), f"'调用链拓扑可视化' 出现了 {all_text.count('调用链拓扑可视化')} 次，预期 1 次"

        print("\n✅ 所有断言通过，无重复文字！")

    finally:
        if os.path.exists(output_path):
            os.remove(output_path)


def test_inline_formatting_preserved():
    """测试粗体、斜体等内联格式在转换后被正确保留"""
    md_content = """# 测试标题

这是一段包含**粗体文本**和*斜体文本*的段落。
"""

    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    temp_file.close()
    output_path = temp_file.name

    try:
        converter = MarkdownConverter()
        converter.to_word(md_content, output_path)

        doc = Document(output_path)

        # 查找包含"粗体文本"的段落
        found_bold = False
        found_italic = False
        for p in doc.paragraphs:
            for run in p.runs:
                if "粗体文本" in run.text and run.bold:
                    found_bold = True
                if "斜体文本" in run.text and run.italic:
                    found_italic = True

        assert found_bold, "粗体格式未被保留"
        assert found_italic, "斜体格式未被保留"

        print("✅ 内联格式（粗体、斜体）正确保留！")

    finally:
        if os.path.exists(output_path):
            os.remove(output_path)


if __name__ == "__main__":
    print("--- 测试 1: 无重复文字 ---")
    test_no_duplicate_text()
    print()
    print("--- 测试 2: 内联格式保留 ---")
    test_inline_formatting_preserved()
    print()
    print("🎉 所有测试通过！")
