#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""使用用户实际文件验证修复效果"""
import io, sys, os

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from src.core.markdown.markdown_converter import MarkdownConverter

bug_dir = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "bug"
)
md_path = os.path.join(bug_dir, "功能汇报.md")
out_path = os.path.join(bug_dir, "功能汇报_fixed.docx")

c = MarkdownConverter()
md = c.read_markdown(md_path)
c.to_word(md, out_path)

doc = Document(out_path)
print("=== 修复后段落 ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i}] {p.text.strip()}")

all_text = " ".join(p.text for p in doc.paragraphs)
print()
print(f"'核心亮点' 出现次数: {all_text.count('核心亮点')}")
print(f"'skywaking' 出现次数: {all_text.count('skywaking')}")
print(f"'功能概述' 出现次数: {all_text.count('功能概述')}")
print(f"'业务系统全景大屏' 出现次数: {all_text.count('业务系统全景大屏')}")
print(f"'调用链拓扑可视化' 出现次数: {all_text.count('调用链拓扑可视化')}")
print(
    f"'异常调用链与慢SQL下钻分析' 出现次数: {all_text.count('异常调用链与慢SQL下钻分析')}"
)
print(f"'请求记录追踪' 出现次数: {all_text.count('请求记录追踪')}")
print(f"'用户轨迹追踪' 出现次数: {all_text.count('用户轨迹追踪')}")

# 所有关键文本应只出现1次
checks = [
    "核心亮点",
    "skywaking",
    "功能概述",
    "业务系统全景大屏",
    "调用链拓扑可视化",
    "异常调用链与慢SQL下钻分析",
    "请求记录追踪",
    "用户轨迹追踪",
]
all_ok = True
for kw in checks:
    cnt = all_text.count(kw)
    if cnt != 1:
        print(f"❌ '{kw}' 出现了 {cnt} 次，预期 1 次")
        all_ok = False

if all_ok:
    print("\n✅ 验证通过！所有关键文本均只出现 1 次，无重复文字。")
else:
    print("\n❌ 仍存在重复文字问题！")
