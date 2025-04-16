#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文件转换器 - Word转换核心模块
版权所有 (c) 2025 Junly
"""

import os
import tempfile
from docx import Document
from docx.shared import Inches
import markdown
import html
from pathlib import Path

class WordConverter:
    """Word文档转换类"""
    
    def __init__(self):
        """初始化Word转换器"""
        pass
    
    def read_word(self, file_path):
        """读取Word文档"""
        try:
            return Document(file_path)
        except Exception as e:
            raise Exception(f"读取Word文档失败: {str(e)}")
    
    def save_as_markdown(self, doc, output_path, include_images=True, progress_callback=None):
        """将Word文档保存为Markdown格式"""
        try:
            md_content = []
            total_paragraphs = len(doc.paragraphs) + sum(1 for table in doc.tables)
            processed = 0
            
            # 处理段落
            for para in doc.paragraphs:
                # 根据样式确定Markdown格式
                if para.style.name.startswith('Heading 1'):
                    md_content.append(f"# {para.text}")
                elif para.style.name.startswith('Heading 2'):
                    md_content.append(f"## {para.text}")
                elif para.style.name.startswith('Heading 3'):
                    md_content.append(f"### {para.text}")
                elif para.style.name.startswith('Heading 4'):
                    md_content.append(f"#### {para.text}")
                elif para.style.name.startswith('Heading 5'):
                    md_content.append(f"##### {para.text}")
                elif para.style.name.startswith('Heading 6'):
                    md_content.append(f"###### {para.text}")
                else:
                    # 处理段落中的图片
                    if include_images and len(para._element.xpath('.//w:drawing')) > 0:
                        # 这里需要提取图片并保存，然后在Markdown中引用
                        # 由于图片处理复杂，这里仅添加图片标记
                        md_content.append(f"{para.text}\n\n![图片描述](image_path)\n")
                    else:
                        md_content.append(f"{para.text}")
                
                # 更新进度
                processed += 1
                if progress_callback:
                    progress_callback(int(processed / total_paragraphs * 100))
            
            # 处理表格
            for table in doc.tables:
                md_table = []
                # 表头
                header_row = []
                for cell in table.rows[0].cells:
                    header_row.append(cell.text.strip())
                md_table.append('| ' + ' | '.join(header_row) + ' |')
                
                # 表头分隔行
                md_table.append('| ' + ' | '.join(['---'] * len(header_row)) + ' |')
                
                # 表格内容
                for row in table.rows[1:]:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    md_table.append('| ' + ' | '.join(row_data) + ' |')
                
                md_content.append('\n'.join(md_table))
                
                # 更新进度
                processed += 1
                if progress_callback:
                    progress_callback(int(processed / total_paragraphs * 100))
            
            # 将内容写入文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n\n'.join(md_content))
            
            return True
            
        except Exception as e:
            raise Exception(f"保存为Markdown失败: {str(e)}")
    
    def save_as_html(self, doc, output_path, include_images=True, progress_callback=None):
        """将Word文档保存为HTML格式"""
        try:
            html_content = ['<!DOCTYPE html>', '<html>', '<head>',
                          '<meta charset="utf-8">',
                          '<title>转换文档</title>',
                          '<style>',
                          'body { font-family: Arial, sans-serif; line-height: 1.6; }',
                          'table { border-collapse: collapse; width: 100%; }',
                          'table, th, td { border: 1px solid #ddd; }',
                          'th, td { padding: 8px; text-align: left; }',
                          'th { background-color: #f2f2f2; }',
                          '</style>',
                          '</head>', '<body>']
            
            total_paragraphs = len(doc.paragraphs) + sum(1 for table in doc.tables)
            processed = 0
            
            # 处理段落
            for para in doc.paragraphs:
                # 根据样式确定HTML标签
                if para.style.name.startswith('Heading 1'):
                    html_content.append(f"<h1>{html.escape(para.text)}</h1>")
                elif para.style.name.startswith('Heading 2'):
                    html_content.append(f"<h2>{html.escape(para.text)}</h2>")
                elif para.style.name.startswith('Heading 3'):
                    html_content.append(f"<h3>{html.escape(para.text)}</h3>")
                elif para.style.name.startswith('Heading 4'):
                    html_content.append(f"<h4>{html.escape(para.text)}</h4>")
                elif para.style.name.startswith('Heading 5'):
                    html_content.append(f"<h5>{html.escape(para.text)}</h5>")
                elif para.style.name.startswith('Heading 6'):
                    html_content.append(f"<h6>{html.escape(para.text)}</h6>")
                else:
                    # 处理段落中的图片
                    if include_images and len(para._element.xpath('.//w:drawing')) > 0:
                        # 这里需要提取图片并保存，然后在HTML中引用
                        # 由于图片处理复杂，这里仅添加图片标记
                        html_content.append(f"<p>{html.escape(para.text)}</p>")
                        html_content.append('<p><img src="image_path" alt="图片描述"></p>')
                    else:
                        html_content.append(f"<p>{html.escape(para.text)}</p>")
                
                # 更新进度
                processed += 1
                if progress_callback:
                    progress_callback(int(processed / total_paragraphs * 100))
            
            # 处理表格
            for table in doc.tables:
                html_table = ['<table>']
                # 表头
                html_table.append('<tr>')
                for cell in table.rows[0].cells:
                    html_table.append(f'<th>{html.escape(cell.text.strip())}</th>')
                html_table.append('</tr>')
                
                # 表格内容
                for row in table.rows[1:]:
                    html_table.append('<tr>')
                    for cell in row.cells:
                        html_table.append(f'<td>{html.escape(cell.text.strip())}</td>')
                    html_table.append('</tr>')
                
                html_table.append('</table>')
                html_content.append('\n'.join(html_table))
                
                # 更新进度
                processed += 1
                if progress_callback:
                    progress_callback(int(processed / total_paragraphs * 100))
            
            html_content.extend(['</body>', '</html>'])
            
            # 将内容写入文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(html_content))
            
            return True
            
        except Exception as e:
            raise Exception(f"保存为HTML失败: {str(e)}")
    
    def remove_images(self, input_path, output_path, progress_callback=None):
        """移除Word文档中的图片后保存为新文档"""
        try:
            doc = self.read_word(input_path)
            
            # 遍历所有段落，查找并移除图片
            total_paragraphs = len(doc.paragraphs)
            for i, para in enumerate(doc.paragraphs):
                try:
                    if len(para._element.xpath('.//w:drawing')) > 0:
                        # 保留段落文本但移除图片元素
                        for run in para.runs:
                            if hasattr(run, '_element'):
                                for drawing in run._element.xpath('.//w:drawing'):
                                    if hasattr(drawing, 'getparent') and drawing.getparent() is not None:
                                        drawing.getparent().remove(drawing)
                except Exception as e:
                    print(f"处理段落 {i} 中的图片时出错: {str(e)}")
                    # 继续处理其他段落，不中断整个流程
                
                # 更新进度
                if progress_callback:
                    progress_callback(int((i+1) / total_paragraphs * 100))
            
            # 保存修改后的文档
            doc.save(output_path)
            return True
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print(f"移除图片失败: {error_details}")
            raise Exception(f"移除图片失败: {str(e)}")
    
    def save_as_text(self, doc, output_path, progress_callback=None):
        """将Word文档保存为纯文本格式"""
        try:
            text_content = []
            total_items = len(doc.paragraphs) + len(doc.tables)
            processed = 0
            
            # 提取所有段落文本
            for i, para in enumerate(doc.paragraphs):
                text_content.append(para.text)
                
                # 更新进度
                processed += 1
                if progress_callback:
                    progress_callback(int(processed / total_items * 100))
            
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join(cell.text for cell in row.cells)
                    text_content.append(row_text)
                
                # 更新进度
                processed += 1
                if progress_callback:
                    progress_callback(int(processed / total_items * 100))
            
            # 将内容写入文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(text_content))
            
            return True
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print(f"保存为文本失败: {error_details}")
            raise Exception(f"保存为文本失败: {str(e)}") 