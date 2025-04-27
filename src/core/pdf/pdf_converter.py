#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Junly文件转换工具 - PDF转换核心模块
版权所有 (c) 2025 Junly
"""

import os
import io
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches
from PIL import Image
import re

class PDFConverter:
    """PDF文档转换类"""
    
    def __init__(self):
        """初始化PDF转换器"""
        pass
    
    def read_pdf(self, file_path):
        """读取PDF文档"""
        try:
            return fitz.open(file_path)
        except Exception as e:
            raise Exception(f"读取PDF文档失败: {str(e)}")
    
    def to_word(self, pdf, output_path, include_images=True, progress_callback=None):
        """将PDF转换为Word文档"""
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置基本样式
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(11)
            
            # 创建图片目录
            if include_images:
                img_dir = os.path.splitext(output_path)[0] + "_images"
                os.makedirs(img_dir, exist_ok=True)
            
            # 处理所有页面
            for i, page in enumerate(pdf):
                # 提取文本
                text = page.get_text()
                
                # 添加页面文本到Word文档
                doc.add_paragraph(text)
                
                # 提取图片
                if include_images:
                    image_list = page.get_images(full=True)
                    for j, img in enumerate(image_list):
                        xref = img[0]
                        try:
                            # 获取图像数据
                            base_image = pdf.extract_image(xref)
                            image_bytes = base_image["image"]
                            
                            # 确定图像类型
                            ext = base_image["ext"]
                            
                            # 保存图像到文件
                            img_path = os.path.join(img_dir, f"page{i+1}_img{j+1}.{ext}")
                            with open(img_path, "wb") as f:
                                f.write(image_bytes)
                            
                            # 将图像添加到Word文档
                            doc.add_picture(img_path, width=Inches(4))
                        except Exception as e:
                            # 图像提取失败，继续处理其他图像
                            print(f"图像提取失败: {str(e)}")
                
                # 更新进度
                if progress_callback:
                    progress_callback(int((i+1) / len(pdf) * 100))
            
            # 保存文档
            doc.save(output_path)
            
            return True
            
        except Exception as e:
            raise Exception(f"转换为Word失败: {str(e)}")
    
    def to_text(self, pdf, output_path, progress_callback=None):
        """将PDF转换为文本文件"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                for i, page in enumerate(pdf):
                    # 提取文本
                    text = page.get_text()
                    f.write(text)
                    
                    # 不同页面之间添加分隔线
                    if i < len(pdf) - 1:
                        f.write("\n\n" + "-" * 80 + "\n\n")
                    
                    # 更新进度
                    if progress_callback:
                        progress_callback(int((i+1) / len(pdf) * 100))
            
            return True
            
        except Exception as e:
            raise Exception(f"转换为文本失败: {str(e)}")
    
    def extract_images(self, pdf, output_dir, progress_callback=None):
        """提取PDF中的图片"""
        try:
            # 确保输出目录存在
            os.makedirs(output_dir, exist_ok=True)
            
            image_count = 0
            
            # 处理所有页面
            for i, page in enumerate(pdf):
                # 提取图片
                image_list = page.get_images(full=True)
                
                for j, img in enumerate(image_list):
                    xref = img[0]
                    try:
                        # 获取图像数据
                        base_image = pdf.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        # 确定图像类型
                        ext = base_image["ext"]
                        
                        # 保存图像到文件
                        img_path = os.path.join(output_dir, f"page{i+1}_img{j+1}.{ext}")
                        with open(img_path, "wb") as f:
                            f.write(image_bytes)
                            
                        image_count += 1
                    except Exception as e:
                        # 图像提取失败，继续处理其他图像
                        print(f"图像提取失败: {str(e)}")
                
                # 更新进度
                if progress_callback:
                    progress_callback(int((i+1) / len(pdf) * 100))
            
            return image_count
            
        except Exception as e:
            raise Exception(f"提取图片失败: {str(e)}") 