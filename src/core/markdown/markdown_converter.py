#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Junly文件工具 - Markdown转换核心模块
版权所有 (c) 2025 Junly
"""

import os
import re
import markdown
import openpyxl
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from bs4 import BeautifulSoup

class MarkdownConverter:
    """Markdown文档转换类"""
    
    def __init__(self):
        """初始化Markdown转换器"""
        pass
    
    def read_markdown(self, file_path):
        """读取Markdown文档"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return content
        except Exception as e:
            raise Exception(f"读取Markdown文档失败: {str(e)}")
    
    def to_word(self, md_content, output_path, progress_callback=None):
        """将Markdown转换为Word文档"""
        try:
            # 将Markdown转换为HTML
            html = markdown.markdown(
                md_content,
                extensions=[
                    'markdown.extensions.tables',
                    'markdown.extensions.fenced_code',
                    'markdown.extensions.codehilite',
                    'markdown.extensions.toc'
                ]
            )
            
            # 使用BeautifulSoup解析HTML
            soup = BeautifulSoup(html, 'html.parser')
            
            # 创建Word文档
            doc = Document()
            
            # 设置基本样式
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(11)
            
            # 总元素数，用于计算进度
            total_elements = len(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table', 'pre']))
            processed = 0
            
            # 处理所有元素
            self._process_html_elements(soup, doc)
            
                # 更新进度
            if progress_callback:
                processed += 1
                progress_callback(int(processed / total_elements * 100))
            
            # 保存文档
            doc.save(output_path)
            
            return True
            
        except Exception as e:
            raise Exception(f"转换为Word失败: {str(e)}")
    
    def _process_html_elements(self, soup, doc):
        """处理HTML元素转换为Word元素"""
        for element in soup.children:
            if element.name is None:
                continue
                
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                p = doc.add_heading('', level=level)
                p.add_run(element.get_text()).bold = True
                
            elif element.name == 'p':
                p = doc.add_paragraph()
                p.add_run(element.get_text())
                
            elif element.name == 'ul' or element.name == 'ol':
                self._process_list(element, doc)
                
            elif element.name == 'table':
                self._process_table(element, doc)
                
            elif element.name == 'pre':
                p = doc.add_paragraph()
                p.add_run(element.get_text()).font.name = 'Courier New'
                
            # 递归处理子元素
            self._process_html_elements(element, doc)
    
    def _process_list(self, list_elem, doc):
        """处理列表元素"""
        is_ordered = list_elem.name == 'ol'
        
        for i, item in enumerate(list_elem.find_all('li', recursive=False)):
            if is_ordered:
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{i+1}. {item.get_text()}")
            else:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"• {item.get_text()}")
    
    def _process_table(self, table_elem, doc):
        """处理表格元素"""
        rows = table_elem.find_all('tr')
        if not rows:
            return
            
        # 创建表格
        num_rows = len(rows)
        max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
        
        if num_rows == 0 or max_cols == 0:
            return
            
        table = doc.add_table(rows=num_rows, cols=max_cols)
        table.style = 'Table Grid'
        
        # 填充表格数据
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                if j < max_cols:  # 防止越界
                    text = cell.get_text().strip()
                    table.cell(i, j).text = text
                    
                    # 如果是表头，应用粗体样式
                    if cell.name == 'th' or i == 0:
                        for paragraph in table.cell(i, j).paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
    
    def extract_tables_to_excel(self, md_content, output_path, progress_callback=None):
        """提取Markdown表格数据，保存为Excel"""
        try:
            # 使用正则表达式匹配Markdown表格
            # 表格模式: |列1|列2|...| + 分隔行 + 数据行
            table_pattern = r'(\|[^\n]+\|\n\|[-:| ]+\|\n(?:\|[^\n]+\|\n)+)'
            tables = re.findall(table_pattern, md_content)
            
            if not tables:
                raise Exception("未找到表格数据")
            
            # 创建Excel工作簿
            workbook = openpyxl.Workbook()
            # 删除默认工作表
            if 'Sheet' in workbook.sheetnames:
                std_sheet = workbook['Sheet']
                workbook.remove(std_sheet)
            
            # 处理找到的表格
            same_header_tables = []
            diff_header_tables = []
            
            for i, table_text in enumerate(tables):
                # 解析表格
                lines = table_text.strip().split('\n')
                header_line = lines[0]
                
                # 提取表头
                headers = [cell.strip() for cell in header_line.strip('|').split('|')]
                
                # 创建表格数据
                table_data = []
                table_data.append(headers)  # 添加表头
                
                # 添加数据行 (跳过分隔行)
                for line in lines[2:]:
                    if line.strip():
                        cells = [cell.strip() for cell in line.strip('|').split('|')]
                        table_data.append(cells)
                
                # 根据表头判断归类
                if i == 0:
                    same_header_tables.append(table_data)
                else:
                    # 与第一个表格比较表头
                    if headers == same_header_tables[0][0]:
                        same_header_tables.append(table_data)
                    else:
                        diff_header_tables.append(table_data)
                
                # 更新进度
                if progress_callback:
                    progress_callback(int((i+1) / len(tables) * 100))
            
            # 创建包含相同表头的工作表
            if same_header_tables:
                same_sheet = workbook.create_sheet(title="相同表头表格")
                row_idx = 1
                
                # 写入表头
                headers = same_header_tables[0][0]
                for col_idx, header in enumerate(headers, 1):
                    same_sheet.cell(row=row_idx, column=col_idx, value=header)
                
                row_idx += 1
                
                # 写入所有表格数据
                for table in same_header_tables:
                    for data_row in table[1:]:  # 跳过表头
                        for col_idx, cell_value in enumerate(data_row, 1):
                            if col_idx <= len(headers):
                                same_sheet.cell(row=row_idx, column=col_idx, value=cell_value)
                        row_idx += 1
            
            # 创建包含不同表头的工作表
            if diff_header_tables:
                diff_sheet = workbook.create_sheet(title="不同表头表格")
                row_idx = 1
                
                # 依次写入每个表格
                for table_idx, table in enumerate(diff_header_tables):
                    # 表格之间空行分隔
                    if table_idx > 0:
                        row_idx += 2
                    
                    # 写入表格内容
                    for data_row in table:
                        for col_idx, cell_value in enumerate(data_row, 1):
                            diff_sheet.cell(row=row_idx, column=col_idx, value=cell_value)
                        row_idx += 1
            
            # 如果没有不同表头的表格，创建空白工作表
            if not diff_header_tables and not same_header_tables:
                workbook.create_sheet(title="无表格数据")
            
            # 保存Excel文件
            workbook.save(output_path)
            
            return True
            
        except Exception as e:
            raise Exception(f"提取表格失败: {str(e)}") 